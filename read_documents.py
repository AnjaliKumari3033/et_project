"""
NovaChem Knowledge Graph - Document Reader (Team B / Person 1)
=================================================================
Reads every .docx and .pdf in the Dataset as ONE full-text document each
(no chunking) so Qwen can extract entities/relationships without losing
context across chunk boundaries.

.xlsx files are intentionally skipped here — equipment_master.xlsx and
equipment_relationships.xlsx are already structured data and should go
straight to Person 2 as a seed graph, not through LLM extraction.

Usage:
    python read_documents.py --dataset_root /path/to/Dataset --out documents.jsonl

Requires: python-docx, pymupdf
    pip install python-docx pymupdf
"""

import argparse
import json
import re
from pathlib import Path

import docx
import fitz  # PyMuPDF

EVENT_ID_RE = re.compile(r"EVT-\d+")
MAX_PDF_PAGES = 60  # safety cap for oversized image-heavy brochures

# Fields worth pulling out of docx tables as ready-made metadata for Qwen
WANTED_FIELDS = (
    "Document UUID", "Document Number", "Event ID", "Chain ID",
    "Equipment ID", "Issue Date", "Revision", "Prepared By", "Approved By",
)


def classify_source(path: Path, dataset_root: Path) -> dict:
    rel = path.relative_to(dataset_root)
    parts = rel.parts

    if parts[0] == "generated_documents":
        doc_type = parts[1] if len(parts) > 1 else "unknown"
        layer = "operational_history"
    elif parts[0] == "public_document":
        doc_type = parts[1] if len(parts) > 1 else "unknown"
        layer = "public_industrial_knowledge"
    else:
        doc_type = path.stem
        layer = "company_foundation"

    event_match = EVENT_ID_RE.search(path.stem)
    event_id = event_match.group(0) if event_match else None

    return {
        "layer": layer,
        "doc_type": doc_type,
        "event_id": event_id,
        "file_path": str(rel),
        "file_name": path.name,
    }


def read_docx(path: Path) -> tuple[str, dict]:
    d = docx.Document(str(path))
    body_lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    body_text = "\n".join(body_lines)

    fields = {}
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) == 2 and cells[0] in WANTED_FIELDS and cells[1]:
                fields[cells[0]] = cells[1]

    return body_text, fields


def read_pdf(path: Path) -> str:
    text_parts = []
    doc = fitz.open(str(path))
    try:
        n_pages = min(len(doc), MAX_PDF_PAGES)
        for i in range(n_pages):
            page_text = doc[i].get_text()
            if page_text:
                text_parts.append(page_text)
    finally:
        doc.close()
    return "\n".join(text_parts)


def process_file(path: Path, dataset_root: Path) -> dict | None:
    meta = classify_source(path, dataset_root)
    suffix = path.suffix.lower()

    try:
        if suffix == ".docx":
            body_text, fields = read_docx(path)
            if not body_text.strip():
                return None
            return {"text": body_text, **meta, **fields}

        elif suffix == ".pdf":
            full_text = read_pdf(path)
            if not full_text.strip():
                return None
            return {"text": full_text, **meta}

    except Exception as e:
        print(f"  [WARN] failed to read {path}: {e}")
        return None

    return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--dataset_root", required=True, help="Path to the Dataset folder")
    parser.add_argument("--out", default="documents.jsonl", help="Output JSONL path")
    args = parser.parse_args()

    dataset_root = Path(args.dataset_root).resolve()
    documents = []
    file_count = 0

    for path in sorted(dataset_root.rglob("*")):
        if path.is_file() and path.suffix.lower() in (".docx", ".pdf"):
            file_count += 1
            doc = process_file(path, dataset_root)
            if doc:
                documents.append(doc)
            if file_count % 100 == 0:
                print(f"  processed {file_count} files, {len(documents)} documents so far...")

    with open(args.out, "w", encoding="utf-8") as f:
        for d in documents:
            f.write(json.dumps(d, default=str, ensure_ascii=False) + "\n")

    print(f"\nDone. {file_count} files -> {len(documents)} documents written to {args.out}")

    from collections import Counter
    doc_type_counts = Counter(d["doc_type"] for d in documents)
    print("\nBy doc_type:")
    for k, v in doc_type_counts.most_common(20):
        print(f"  {k}: {v} documents")


if __name__ == "__main__":
    main()
