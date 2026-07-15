"""
=========================================================
NovaChem Industrial Knowledge Intelligence
Document Format Converter
=========================================================
"""

import os
import shutil

from docx2pdf import convert
from docx import Document
from openpyxl import Workbook

from config import *

# ==========================================================
# CONVERT DOCX -> TXT
# ==========================================================

def docx_to_txt(docx_file):

    txt_file = docx_file.replace(".docx", ".txt")

    doc = Document(docx_file)

    with open(txt_file, "w", encoding="utf-8") as f:

        for para in doc.paragraphs:
            f.write(para.text + "\n")

    print(f"[TXT] {os.path.basename(txt_file)}")


# ==========================================================
# CONVERT DOCX -> XLSX
# ==========================================================

def docx_to_xlsx(docx_file):

    xlsx_file = docx_file.replace(".docx", ".xlsx")

    wb = Workbook()

    ws = wb.active

    ws.title = "Document"

    doc = Document(docx_file)

    row = 1

    for para in doc.paragraphs:

        text = para.text.strip()

        if text:

            ws.cell(row=row, column=1).value = text
            row += 1

    wb.save(xlsx_file)

    print(f"[XLSX] {os.path.basename(xlsx_file)}")


# ==========================================================
# CONVERT DOCX -> PDF
# ==========================================================

def docx_to_pdf(docx_file):

    convert(docx_file)

    pdf_file = docx_file.replace(".docx", ".pdf")

    print(f"[PDF] {os.path.basename(pdf_file)}")


# ==========================================================
# PROCESS DIRECTORY
# ==========================================================

def process_folder(folder, output_type):

    if not os.path.exists(folder):
        return

    for file in os.listdir(folder):

        if not file.endswith(".docx"):
            continue

        path = os.path.join(folder, file)

        try:

            if output_type == "pdf":
                docx_to_pdf(path)

            elif output_type == "txt":
                docx_to_txt(path)

            elif output_type == "xlsx":
                docx_to_xlsx(path)

            elif output_type == "docx":
                pass

        except Exception as e:

            print(f"[ERROR] {file}")
            print(e)


# ==========================================================
# MAIN
# ==========================================================

print("\nConverting Generated Documents...\n")

process_folder(MAINTENANCE_DIR, "docx")
process_folder(INSPECTION_DIR, "pdf")
process_folder(WORKORDER_DIR, "xlsx")
process_folder(INCIDENT_DIR, "pdf")
process_folder(EMAIL_DIR, "txt")
process_folder(RCA_DIR, "docx")
process_folder(QUALITY_DIR, "xlsx")
process_folder(SAFETY_DIR, "pdf")
process_folder(CALIBRATION_DIR, "pdf")

print("\nDone!\n")