"""
=========================================================
NovaChem Industrial Knowledge Intelligence
Document Utility Functions
=========================================================
"""

import random
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config import *


# ==========================================================
# CREATE DOCUMENT
# ==========================================================

def create_document(title):

    doc = Document()

    section = doc.sections[0]

    section.left_margin = Inches(PAGE_MARGIN)
    section.right_margin = Inches(PAGE_MARGIN)
    section.top_margin = Inches(PAGE_MARGIN)
    section.bottom_margin = Inches(PAGE_MARGIN)

    style = doc.styles["Normal"]

    style.font.name = DEFAULT_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    style.font.size = Pt(FONT_SIZE)

    add_company_header(doc, title)

    return doc


# ==========================================================
# COMPANY HEADER
# ==========================================================

def add_company_header(doc, title):

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run(COMPANY_NAME + "\n")
    run.bold = True
    run.font.size = Pt(TITLE_SIZE)

    run = p.add_run(PLANT_NAME + "\n")
    run.bold = True
    run.font.size = Pt(14)

    run = p.add_run(COMPANY_ADDRESS + "\n\n")
    run.font.size = Pt(10)

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()


# ==========================================================
# DOCUMENT METADATA
# ==========================================================

def add_document_metadata(doc, event, document_number):

    doc.add_heading("Document Information", level=1)

    table = doc.add_table(rows=8, cols=2)

    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    data = [

        ("Document UUID", str(uuid.uuid4())[:8].upper()),

        ("Document Number", document_number),

        ("Event ID", event["Event ID"]),

        ("Chain ID", event["Chain_ID"]),

        ("Equipment ID", event["Equipment ID"]),

        ("Department", event["Department"]),

        ("Generated On",
         datetime.now().strftime("%d-%b-%Y %H:%M")),

        ("Confidentiality",
         CONFIDENTIALITY)

    ]

    for row, item in zip(table.rows, data):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()


# ==========================================================
# DOCUMENT CONTROL
# ==========================================================

def add_document_control(
        doc,
        document_number,
        issue_date,
        prepared_by,
        approved_by,
        revision
):

    doc.add_heading("Document Control", level=1)

    table = doc.add_table(rows=6, cols=2)

    table.style = "Table Grid"

    values = [

        ("Document Number", document_number),

        ("Issue Date", str(issue_date)),

        ("Revision", revision),

        ("Prepared By", prepared_by),

        ("Approved By", approved_by),

        ("Status", "Approved")

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = item[0]

        row.cells[1].text = str(item[1])

    doc.add_paragraph()


# ==========================================================
# EQUIPMENT DETAILS
# ==========================================================

def add_equipment_details(doc, event):

    doc.add_heading("Equipment Details", level=1)

    table = doc.add_table(rows=9, cols=2)

    table.style = "Table Grid"

    values = [

        ("Equipment ID", event["Equipment ID"]),

        ("Equipment Name", event["Equipment Name"]),

        ("Area", event["Area"]),

        ("Department", event["Department"]),

        ("Shift", event["Shift"]),

        ("Reported From", event["Reported_From"]),

        ("Asset Criticality",
         event["Asset Criticality"]),

        ("Failure Category",
         event["Failure Category"]),

        ("Downtime",
         f"{event['Downtime (hrs)']} Hours")

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = item[0]

        row.cells[1].text = str(item[1])

    doc.add_paragraph()


# ==========================================================
# SECTION HEADING
# ==========================================================

def add_section(doc, title):

    doc.add_heading(title, level=1)


# ==========================================================
# PARAGRAPH
# ==========================================================

def add_text(doc, text):

    doc.add_paragraph(str(text))


# ==========================================================
# SUMMARY TABLE
# ==========================================================

# ==========================================================
# SUMMARY TABLE
# ==========================================================

def add_summary_table(doc, event, title="Maintenance Summary"):

    doc.add_heading(title, level=1)

    table = doc.add_table(rows=6, cols=2)

    table.style = "Table Grid"

    values = [

        ("Priority", event["Priority"]),

        ("Severity", event["Severity"]),

        ("Status", event["Status"]),

        ("Downtime",
         f"{event['Downtime (hrs)']} Hours"),

        ("Follow-up Required",
         event["Follow-up Required"]),

        ("Asset Criticality",
         event["Asset Criticality"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = item[0]

        row.cells[1].text = str(item[1])

    doc.add_paragraph()


# ==========================================================
# SPARE PARTS TABLE
# ==========================================================

def add_spare_parts(doc, event):

    doc.add_heading("Spare Parts Used", level=1)

    table = doc.add_table(rows=4, cols=2)

    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Part Name"
    table.rows[0].cells[1].text = "Quantity"

    category = str(event["Failure Category"]).lower()

    if "mechanical" in category:

        parts = [
            ("Bearing", "1"),
            ("Mechanical Seal", "1"),
            ("Grease Cartridge", "2")
        ]

    elif "electrical" in category:

        parts = [
            ("Fuse", "2"),
            ("Terminal Lug", "4"),
            ("Power Cable", "2 m")
        ]

    elif "instrument" in category:

        parts = [
            ("Pressure Transmitter", "1"),
            ("Signal Cable", "3 m"),
            ("Calibration Kit", "1")
        ]

    else:

        parts = [
            ("Standard Spare", "1"),
            ("Lubricant", "1"),
            ("Fastener Kit", "1")
        ]

    for i, part in enumerate(parts):

        table.rows[i + 1].cells[0].text = part[0]

        table.rows[i + 1].cells[1].text = part[1]

    doc.add_paragraph()


# ==========================================================
# CHECKLIST
# ==========================================================

def add_checklist(
    doc,
    allow_variation=True,
    title="Maintenance Checklist"
):

    doc.add_heading(title, level=1)

    checklist = [
        "Lock Out / Tag Out (LOTO) Followed",
        "PPE Worn",
        "Visual Inspection Completed",
        "Lubrication Checked",
        "Leak Inspection Completed",
        "Alignment Verified",
        "Trial Run Successful"
    ]

    skip_index = None

    if allow_variation and random.random() < 0.15:

        skip_index = random.randrange(len(checklist))

    for i, item in enumerate(checklist):

        if i == skip_index:

            doc.add_paragraph(f"☐ {item} — Pending")

        else:

            doc.add_paragraph(f"☑ {item}")

    doc.add_paragraph()


# ==========================================================
# RECOMMENDATION
# ==========================================================

def add_recommendation(doc, text):

    doc.add_heading("Recommendations", level=1)

    doc.add_paragraph(str(text))

    doc.add_paragraph()


# ==========================================================
# LINKED DOCUMENTS
# ==========================================================

def add_linked_documents(doc, docs):

    doc.add_heading("Linked Documents", level=1)

    if len(docs) == 0:

        doc.add_paragraph("No linked documents for this event.")

        return

    for d in docs:

        doc.add_paragraph(
            d,
            style="List Bullet"
        )

    doc.add_paragraph()


# ==========================================================
# APPROVAL
# ==========================================================

def add_signature_block(doc, prepared_by, approved_by):

    doc.add_heading("Approval", level=1)

    table = doc.add_table(rows=2, cols=2)

    table.style = "Table Grid"

    table.rows[0].cells[0].text = "Prepared By"

    table.rows[0].cells[1].text = str(prepared_by)

    table.rows[1].cells[0].text = "Approved By"

    table.rows[1].cells[1].text = str(approved_by)

    doc.add_paragraph()


# ==========================================================
# FOOTER
# ==========================================================

def add_footer(doc):

    footer = doc.sections[0].footer

    para = footer.paragraphs[0]

    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    para.text = (
        f"{COMPANY_NAME} | "
        f"{CONFIDENTIALITY} | "
        f"Generated by Industrial Knowledge Intelligence"
    )


# ==========================================================
# SAVE DOCUMENT
# ==========================================================

def save_document(doc, filepath):

    add_footer(doc)

    doc.save(filepath)


# ==========================================================
# SIMPLE KEY-VALUE TABLE
# ==========================================================

def add_key_value_table(doc, heading, values):

    doc.add_heading(heading, level=1)

    table = doc.add_table(
        rows=len(values),
        cols=2
    )

    table.style = "Table Grid"

    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()


# ==========================================================
# MAINTENANCE HISTORY
# ==========================================================
# Expects history_df to already be filtered to events on/before
# the current document's date — see utils.equipment_history().

def add_history(doc, history_df):

    doc.add_heading(
        "Previous Equipment History",
        level=1
    )

    if len(history_df) == 0:

        doc.add_paragraph(
            "No previous maintenance history available."
        )

        return

    table = doc.add_table(rows=1, cols=4)

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "Event"
    hdr[1].text = "Date"
    hdr[2].text = "Failure"
    hdr[3].text = "Status"

    for _, row in history_df.iterrows():

        cells = table.add_row().cells

        cells[0].text = str(row["Event ID"])

        cells[1].text = str(row["Date"])

        cells[2].text = str(row["Root Cause"])

        cells[3].text = str(row["Status"])

    doc.add_paragraph()

# ==========================================================
# WORK ORDER INFORMATION
# ==========================================================

def add_workorder_information(doc, event):

    doc.add_heading("Work Order Information", level=1)

    table = doc.add_table(rows=8, cols=2)

    table.style = "Table Grid"

    values = [

        ("Work Order Number", f"WO_{event['Event ID'].replace('EVT', '')}"),

        ("Event ID", event["Event ID"]),

        ("Assigned To", event["Assigned To"]),

        ("Department", event["Department"]),

        ("Priority", event["Priority"]),

        ("Status", event["Status"]),

        ("Shift", event["Shift"]),

        ("Planned Duration",
         f"{event['Downtime (hrs)']} Hours")

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = item[0]

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

# ==========================================================
# REQUIRED RESOURCES
# ==========================================================

def add_required_resources(doc, event, tool):

    doc.add_heading("Required Resources", level=1)

    table = doc.add_table(rows=4, cols=2)

    table.style = "Table Grid"

    values = [

        ("Assigned Technician",
         event["Assigned To"]),

        ("Required Tool",
         tool),

        ("Department",
         event["Department"]),

        ("Failure Category",
         event["Failure Category"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = item[0]

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

# ==========================================================
# ESTIMATED DURATION
# ==========================================================

def add_estimated_duration(doc, duration):

    doc.add_heading("Estimated Duration", level=1)

    doc.add_paragraph(str(duration))

    doc.add_paragraph()

# ==========================================================
# SAFETY REQUIREMENTS
# ==========================================================

def add_safety_requirements(doc, text):

    doc.add_heading("Safety Requirements", level=1)

    doc.add_paragraph(str(text))

    doc.add_paragraph()

# ==========================================================
# SPECIAL INSTRUCTIONS
# ==========================================================

def add_special_instructions(doc, text):

    doc.add_heading("Special Instructions", level=1)

    doc.add_paragraph(str(text))

    doc.add_paragraph()

# ==========================================================
# INCIDENT INFORMATION
# ==========================================================

def add_incident_information(doc, event):

    doc.add_heading("Incident Information", level=1)

    table = doc.add_table(rows=8, cols=2)

    table.style = "Table Grid"

    values = [

        ("Incident Number", f"INC_{event['Event ID'].replace('EVT', '')}"),

        ("Event ID", event["Event ID"]),

        ("Date", str(event["Date"])),

        ("Department", event["Department"]),

        ("Area", event["Area"]),

        ("Shift", event["Shift"]),

        ("Reported By", event["Reported_From"]),

        ("Status", event["Status"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

# ==========================================================
# IMMEDIATE ACTION
# ==========================================================

def add_immediate_action(doc, text):

    doc.add_heading("Immediate Action Taken", level=1)

    doc.add_paragraph(str(text))

    doc.add_paragraph()

# ==========================================================
# IMPACT ASSESSMENT
# ==========================================================

def add_impact_assessment(doc, event):

    doc.add_heading("Impact Assessment", level=1)

    table = doc.add_table(rows=4, cols=2)

    table.style = "Table Grid"

    values = [

        ("Severity", event["Severity"]),

        ("Priority", event["Priority"]),

        ("Downtime", f"{event['Downtime (hrs)']} Hours"),

        ("Asset Criticality", event["Asset Criticality"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

# ==========================================================
# ROOT CAUSE
# ==========================================================

def add_root_cause(doc, event):

    doc.add_heading("Root Cause", level=1)

    doc.add_paragraph(str(event["Root Cause"]))

    doc.add_paragraph()

# ==========================================================
# SAFETY CLASSIFICATION
# ==========================================================

def add_safety_classification(doc, event):

    doc.add_heading("Safety Classification", level=1)

    table = doc.add_table(rows=3, cols=2)

    table.style = "Table Grid"

    values = [

        ("Failure Category", event["Failure Category"]),

        ("Severity", event["Severity"]),

        ("Follow-up Required", event["Follow-up Required"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

    # ==========================================================
# EMAIL INFORMATION
# ==========================================================

def add_email_information(doc, event, template):

    doc.add_heading("Email Information", level=1)

    table = doc.add_table(rows=5, cols=2)

    table.style = "Table Grid"

    values = [

        ("To", f"{event['Department']} Department"),

        ("From", event["Assigned To"]),

        ("Subject", template["subject"]),

        ("Date", str(event["Date"])),

        ("Related Event", event["Event ID"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

# ==========================================================
# RCA INFORMATION
# ==========================================================

def add_rca_information(doc, event):

    doc.add_heading("RCA Information", level=1)

    table = doc.add_table(rows=7, cols=2)

    table.style = "Table Grid"

    values = [

        ("RCA Number",
         f"RCA_{event['Event ID'].replace('EVT', '')}"),

        ("Event ID",
         event["Event ID"]),

        ("Department",
         event["Department"]),

        ("Reported By",
         event["Reported_From"]),

        ("Assigned To",
         event["Assigned To"]),

        ("Priority",
         event["Priority"]),

        ("Status",
         event["Status"])

    ]

    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])

    doc.add_paragraph()

    # ==========================================================
# QUALITY INFORMATION
# ==========================================================

def add_quality_information(doc, event):

    doc.add_heading(
        "Quality Information",
        level=1
    )

    table = doc.add_table(
        rows=7,
        cols=2
    )

    table.style = "Table Grid"


    values = [

        ("Quality Report Number",
         f"QA_{event['Event ID'].replace('EVT','')}"),

        ("Event ID",
         event["Event ID"]),

        ("Department",
         event["Department"]),

        ("Equipment",
         event["Equipment Name"]),

        ("Severity",
         event["Severity"]),

        ("Priority",
         event["Priority"]),

        ("Status",
         event["Status"])

    ]


    for row,item in zip(table.rows,values):

        row.cells[0].text = item[0]

        row.cells[1].text = str(item[1])


    doc.add_paragraph()



# ==========================================================
# QUALITY ASSESSMENT
# ==========================================================

def add_quality_assessment(doc,event):

    doc.add_heading(
        "Quality Assessment",
        level=1
    )


    table = doc.add_table(
        rows=5,
        cols=2
    )


    table.style="Table Grid"


    values=[

        ("Failure Category",
         event["Failure Category"]),


        ("Root Cause",
         event["Root Cause"]),


        ("Corrective Action",
         event["Corrective Action"]),


        ("Follow-up Required",
         event["Follow-up Required"]),


        ("Asset Criticality",
         event["Asset Criticality"])

    ]


    for row,item in zip(table.rows,values):

        row.cells[0].text=item[0]

        row.cells[1].text=str(item[1])


    doc.add_paragraph()

# ==========================================================
# SAFETY INFORMATION
# ==========================================================

def add_safety_information(doc, event):

    doc.add_heading(
        "Safety Information",
        level=1
    )


    table = doc.add_table(
        rows=8,
        cols=2
    )


    table.style = "Table Grid"


    values = [

        ("Safety Report Number",
         f"SAF_{event['Event ID'].replace('EVT','')}"),


        ("Event ID",
         event["Event ID"]),


        ("Department",
         event["Department"]),


        ("Area",
         event["Area"]),


        ("Shift",
         event["Shift"]),


        ("Reported By",
         event["Reported_From"]),


        ("Severity",
         event["Severity"]),


        ("Status",
         event["Status"])

    ]


    for row, item in zip(table.rows, values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])


    doc.add_paragraph()

# ==========================================================
# CALIBRATION INFORMATION
# ==========================================================

def add_calibration_information(doc, event):

    doc.add_heading(
        "Calibration Information",
        level=1
    )


    table = doc.add_table(
        rows=7,
        cols=2
    )

    table.style = "Table Grid"


    values = [

        (
            "Calibration Number",
            f"CAL_{event['Event ID'].replace('EVT','')}"
        ),

        (
            "Event ID",
            event["Event ID"]
        ),

        (
            "Equipment ID",
            event["Equipment ID"]
        ),

        (
            "Instrument Name",
            event["Equipment Name"]
        ),

        (
            "Department",
            event["Department"]
        ),

        (
            "Assigned Technician",
            event["Assigned To"]
        ),

        (
            "Status",
            event["Status"]
        )

    ]


    for row,item in zip(table.rows,values):

        row.cells[0].text = str(item[0])

        row.cells[1].text = str(item[1])


    doc.add_paragraph()



# ==========================================================
# CALIBRATION RESULT
# ==========================================================

def add_calibration_result(doc, template):

    doc.add_heading(
        "Calibration Result",
        level=1
    )


    add_text(
        doc,
        template["result"]
    )



# ==========================================================
# CALIBRATION FINDINGS
# ==========================================================

def add_calibration_findings(doc, template):

    doc.add_heading(
        "Calibration Findings",
        level=1
    )


    add_text(
        doc,
        template["finding"]
    )



# ==========================================================
# CALIBRATION EQUIPMENT DETAILS
# ==========================================================

def add_calibration_equipment(doc, template):

    doc.add_heading(
        "Calibration Equipment",
        level=1
    )


    add_text(
        doc,
        template["equipment"]
    )