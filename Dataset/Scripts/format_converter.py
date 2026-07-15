"""
=========================================================
NovaChem
Format Conversion Pipeline
=========================================================
"""

import os
from pathlib import Path

from docx import Document
import pandas as pd

from docx2pdf import convert


# =====================================================
# DOCX -> TXT
# =====================================================

def docx_to_txt(docx_file, txt_file):

    doc = Document(docx_file)

    with open(
        txt_file,
        "w",
        encoding="utf-8"
    ) as f:

        for para in doc.paragraphs:

            f.write(para.text + "\n")


# =====================================================
# DOCX -> PDF
# =====================================================

def docx_to_pdf(docx_file, pdf_file):

    convert(docx_file, pdf_file)


# =====================================================
# DOCX -> XLSX
# =====================================================

def docx_to_xlsx(docx_file, xlsx_file):

    doc = Document(docx_file)

    rows = []

    for para in doc.paragraphs:

        text = para.text.strip()

        if text:

            rows.append([text])

    df = pd.DataFrame(rows)

    df.to_excel(
        xlsx_file,
        index=False,
        header=False
    )